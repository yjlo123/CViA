from cvConverter.BaseConverter import BaseConverter
import docx

__author__ = 'siwei'

class DOCXFileConverter(BaseConverter):

    def __init__(self):
        self.type = "docx"

    def convert(self):
        document = docx.Document(self.path)
        #return document.paragraphs[0].text
        docText = '\n'.join([
            paragraph.text.encode('utf-8') for paragraph in document.paragraphs if paragraph.text != ""
        ])
        return self.text+docText