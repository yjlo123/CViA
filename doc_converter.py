__author__ = 'siwei'

from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import TextConverter
from pdfminer.layout import LAParams
from pdfminer.pdfpage import PDFPage
from cStringIO import StringIO
from subprocess import Popen, PIPE
import docx

class DocConverter:
    def __init__(self):
        "123"

    def convertFromDocx(self, file_name):
        document = docx.Document(file_name)
        #return document.paragraphs[0].text
        docText = '\n'.join([
            paragraph.text.encode('utf-8') for paragraph in document.paragraphs if paragraph.text != ""
        ])
        return docText

    def convertFromPdf(self, path):
        rsrcmgr = PDFResourceManager()
        retstr = StringIO()
        codec = 'utf-8'
        laparams = LAParams()
        device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
        fp = file(path, 'rb')
        interpreter = PDFPageInterpreter(rsrcmgr, device)
        password = ""
        maxpages = 0
        caching = True
        pagenos=set()
        for page in PDFPage.get_pages(fp, pagenos, maxpages=maxpages, password=password,caching=caching, check_extractable=True):
            interpreter.process_page(page)
        fp.close()
        device.close()
        str = retstr.getvalue()
        retstr.close()
        return str

    def documentToText(self, file_name):
        if file_name[-4:] == ".doc":
            cmd = ['antiword', file_name]
            p = Popen(cmd, stdout=PIPE)
            stdout, stderr = p.communicate()
            return stdout.decode('ascii', 'ignore')
        elif file_name[-5:] == ".docx":
            return self.convertFromDocx(file_name)
        elif file_name[-4:] == ".txt":
            return open(file_name).read()
        elif file_name[-4:] == ".pdf":
            return self.convertFromPdf(file_name)

if __name__ == "__main__":
    x = DocConverter()
    print x.documentToText("/Users/haojiang/Desktop/CViA/cv/DesmondLim.pdf")
